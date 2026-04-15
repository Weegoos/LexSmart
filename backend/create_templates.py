"""
Generate two .docx labor contract templates for Kazakhstan:
  1. labor_contract_ip.docx  — for Individual Entrepreneur (ИП)
  2. labor_contract_too.docx — for Limited Liability Partnership (ТОО)

Both files are written to ./templates/ relative to this script.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────
def set_cell_text(cell, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Write styled text into a single table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_paragraph(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=4):
    """Append a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_empty_line(doc):
    add_paragraph(doc, "", size=8, space_before=0, space_after=0)


def set_default_style(doc):
    """Set default font for the whole document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)


# ──────────────────────────────────────────────
# Content builders
# ──────────────────────────────────────────────
def build_header(doc, employer_type_label):
    """Title + date/city line."""
    add_paragraph(doc, "ТРУДОВОЙ ДОГОВОР", bold=True, size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, f"({employer_type_label})", bold=False, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_empty_line(doc)
    add_paragraph(doc, "г. Алматы                                                                              "
                       "{{CURRENT_DATE}}",
                  size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_empty_line(doc)


def build_parties_ip(doc):
    """Section 'Стороны' for ИП."""
    add_paragraph(doc, "1. СТОРОНЫ ДОГОВОРА", bold=True, size=12, space_before=6)
    add_paragraph(doc,
        "Работодатель: Индивидуальный предприниматель {{EMPLOYER_NAME}}, "
        "ИИН {{EMPLOYER_IIN_BIN}}, адрес: {{EMPLOYER_ADDRESS}}, "
        "действующий на основании свидетельства о регистрации в качестве "
        "индивидуального предпринимателя (далее — «Работодатель»), с одной стороны, и",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Работник: {{EMPLOYEE_NAME}}, ИИН {{EMPLOYEE_IIN}}, "
        "адрес: {{EMPLOYEE_ADDRESS}} (далее — «Работник»), с другой стороны,",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "совместно именуемые «Стороны», заключили настоящий трудовой договор "
        "(далее — «Договор») о нижеследующем:",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_parties_too(doc):
    """Section 'Стороны' for ТОО."""
    add_paragraph(doc, "1. СТОРОНЫ ДОГОВОРА", bold=True, size=12, space_before=6)
    add_paragraph(doc,
        "Работодатель: Товарищество с ограниченной ответственностью «{{EMPLOYER_NAME}}», "
        "БИН {{EMPLOYER_IIN_BIN}}, адрес: {{EMPLOYER_ADDRESS}}, "
        "в лице директора, действующего на основании Устава "
        "(далее — «Работодатель»), с одной стороны, и",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Работник: {{EMPLOYEE_NAME}}, ИИН {{EMPLOYEE_IIN}}, "
        "адрес: {{EMPLOYEE_ADDRESS}} (далее — «Работник»), с другой стороны,",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "совместно именуемые «Стороны», заключили настоящий трудовой договор "
        "(далее — «Договор») о нижеследующем:",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_subject(doc):
    """Section 2 — Subject of contract."""
    add_paragraph(doc, "2. ПРЕДМЕТ ДОГОВОРА", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "2.1. Работодатель принимает Работника на должность: {{POSITION}}.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "2.2. Работа по настоящему Договору является для Работника основным "
        "местом работы.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "2.3. Место выполнения работы: {{EMPLOYER_ADDRESS}}.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_term(doc):
    """Section 3 — Term, start/end, probation."""
    add_paragraph(doc, "3. СРОК ДЕЙСТВИЯ ДОГОВОРА", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "3.1. Настоящий Договор вступает в силу с «{{START_DATE}}» и заключён "
        "на срок до «{{END_DATE}}».",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "3.2. Работнику устанавливается испытательный срок "
        "продолжительностью {{PROBATION_MONTHS}} месяц(ев) с даты начала работы, "
        "в соответствии со статьёй 36 Трудового кодекса Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "3.3. В период испытательного срока на Работника распространяются нормы "
        "настоящего Договора и трудового законодательства Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_working_conditions(doc):
    """Section 4 — Schedule, vacation."""
    add_paragraph(doc, "4. УСЛОВИЯ ТРУДА", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "4.1. Работнику устанавливается режим рабочего времени: {{WORK_SCHEDULE}}.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "4.2. Работнику предоставляется ежегодный оплачиваемый трудовой отпуск "
        "продолжительностью {{VACATION_DAYS}} календарных дней в соответствии со "
        "статьёй 88 Трудового кодекса Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "4.3. Работодатель обеспечивает безопасные условия труда в соответствии "
        "с требованиями законодательства Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_compensation(doc):
    """Section 5 — Salary."""
    add_paragraph(doc, "5. ОПЛАТА ТРУДА", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "5.1. За выполнение обязанностей по настоящему Договору Работнику "
        "устанавливается ежемесячная заработная плата в размере {{SALARY}} "
        "(до удержания налогов и обязательных отчислений).",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "5.2. Заработная плата выплачивается не реже одного раза в месяц, "
        "не позднее 10 числа месяца, следующего за отработанным, путём "
        "перечисления на банковский счёт Работника.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "5.3. Из заработной платы Работника удерживаются обязательные пенсионные "
        "взносы, индивидуальный подоходный налог и иные обязательные платежи "
        "в соответствии с законодательством Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_rights_obligations(doc):
    """Section 6 — Rights & obligations (brief)."""
    add_paragraph(doc, "6. ПРАВА И ОБЯЗАННОСТИ СТОРОН", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "6.1. Права и обязанности Работодателя определяются настоящим Договором "
        "и Трудовым кодексом Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "6.2. Права и обязанности Работника определяются настоящим Договором, "
        "должностной инструкцией и Трудовым кодексом Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "6.3. Работник обязуется добросовестно исполнять свои трудовые обязанности, "
        "соблюдать правила внутреннего трудового распорядка, трудовую дисциплину, "
        "требования по охране труда и обеспечению безопасности труда.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_custom_clauses(doc):
    """Section 7 — Custom / additional clauses placeholder."""
    add_paragraph(doc, "7. ДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "{{CUSTOM_CLAUSES}}",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_termination(doc):
    """Section 8 — Termination."""
    add_paragraph(doc, "8. ПРЕКРАЩЕНИЕ ДОГОВОРА", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "8.1. Настоящий Договор может быть прекращён по основаниям, "
        "предусмотренным Трудовым кодексом Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "8.2. Сторона, инициирующая расторжение Договора, обязана уведомить "
        "другую сторону в письменной форме не менее чем за один месяц, "
        "если иной срок не предусмотрен законодательством.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_final(doc):
    """Section 9 — Final provisions."""
    add_paragraph(doc, "9. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", bold=True, size=12, space_before=8)
    add_paragraph(doc,
        "9.1. Настоящий Договор составлен в двух экземплярах, имеющих одинаковую "
        "юридическую силу, по одному для каждой из Сторон.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "9.2. Все изменения и дополнения к настоящему Договору оформляются "
        "в письменной форме дополнительными соглашениями.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "9.3. Во всём, что не предусмотрено настоящим Договором, Стороны "
        "руководствуются действующим трудовым законодательством Республики Казахстан.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_signatures_ip(doc):
    """Signature block for ИП."""
    add_empty_line(doc)
    add_paragraph(doc, "10. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН", bold=True, size=12, space_before=8)
    add_empty_line(doc)

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Employer column
    set_cell_text(table.cell(0, 0), "РАБОТОДАТЕЛЬ", bold=True, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 0), "ИП {{EMPLOYER_NAME}}", size=11)
    set_cell_text(table.cell(2, 0), "ИИН: {{EMPLOYER_IIN_BIN}}", size=11)
    set_cell_text(table.cell(3, 0), "Адрес: {{EMPLOYER_ADDRESS}}", size=11)
    set_cell_text(table.cell(4, 0), "", size=11)
    set_cell_text(table.cell(5, 0), "___________________ / {{EMPLOYER_NAME}} /", size=11)
    set_cell_text(table.cell(6, 0), "Дата: {{CURRENT_DATE}}", size=11)

    # Employee column
    set_cell_text(table.cell(0, 1), "РАБОТНИК", bold=True, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), "{{EMPLOYEE_NAME}}", size=11)
    set_cell_text(table.cell(2, 1), "ИИН: {{EMPLOYEE_IIN}}", size=11)
    set_cell_text(table.cell(3, 1), "Адрес: {{EMPLOYEE_ADDRESS}}", size=11)
    set_cell_text(table.cell(4, 1), "", size=11)
    set_cell_text(table.cell(5, 1), "___________________ / {{EMPLOYEE_NAME}} /", size=11)
    set_cell_text(table.cell(6, 1), "Дата: {{CURRENT_DATE}}", size=11)


def build_signatures_too(doc):
    """Signature block for ТОО."""
    add_empty_line(doc)
    add_paragraph(doc, "10. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН", bold=True, size=12, space_before=8)
    add_empty_line(doc)

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Employer column
    set_cell_text(table.cell(0, 0), "РАБОТОДАТЕЛЬ", bold=True, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 0), "ТОО «{{EMPLOYER_NAME}}»", size=11)
    set_cell_text(table.cell(2, 0), "БИН: {{EMPLOYER_IIN_BIN}}", size=11)
    set_cell_text(table.cell(3, 0), "Адрес: {{EMPLOYER_ADDRESS}}", size=11)
    set_cell_text(table.cell(4, 0), "", size=11)
    set_cell_text(table.cell(5, 0), "___________________ / {{EMPLOYER_NAME}} /", size=11)
    set_cell_text(table.cell(6, 0), "Дата: {{CURRENT_DATE}}", size=11)

    # Employee column
    set_cell_text(table.cell(0, 1), "РАБОТНИК", bold=True, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), "{{EMPLOYEE_NAME}}", size=11)
    set_cell_text(table.cell(2, 1), "ИИН: {{EMPLOYEE_IIN}}", size=11)
    set_cell_text(table.cell(3, 1), "Адрес: {{EMPLOYEE_ADDRESS}}", size=11)
    set_cell_text(table.cell(4, 1), "", size=11)
    set_cell_text(table.cell(5, 1), "___________________ / {{EMPLOYEE_NAME}} /", size=11)
    set_cell_text(table.cell(6, 1), "Дата: {{CURRENT_DATE}}", size=11)


# ──────────────────────────────────────────────
# Template generators
# ──────────────────────────────────────────────
def create_ip_template(output_path: str):
    """Generate the ИП (Individual Entrepreneur) labor contract template."""
    doc = Document()
    set_default_style(doc)

    build_header(doc, "для Индивидуального Предпринимателя")
    build_parties_ip(doc)
    build_subject(doc)
    build_term(doc)
    build_working_conditions(doc)
    build_compensation(doc)
    build_rights_obligations(doc)
    build_custom_clauses(doc)
    build_termination(doc)
    build_final(doc)
    build_signatures_ip(doc)

    doc.save(output_path)
    print(f"[OK] Created: {output_path}")


def create_too_template(output_path: str):
    """Generate the ТОО (LLP) labor contract template."""
    doc = Document()
    set_default_style(doc)

    build_header(doc, "для Товарищества с ограниченной ответственностью")
    build_parties_too(doc)
    build_subject(doc)
    build_term(doc)
    build_working_conditions(doc)
    build_compensation(doc)
    build_rights_obligations(doc)
    build_custom_clauses(doc)
    build_termination(doc)
    build_final(doc)
    build_signatures_too(doc)

    doc.save(output_path)
    print(f"[OK] Created: {output_path}")


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────
if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    templates_dir = os.path.join(script_dir, "templates")
    os.makedirs(templates_dir, exist_ok=True)

    ip_path = os.path.join(templates_dir, "labor_contract_ip.docx")
    too_path = os.path.join(templates_dir, "labor_contract_too.docx")

    create_ip_template(ip_path)
    create_too_template(too_path)

    print("\nDone. Both templates generated successfully.")
