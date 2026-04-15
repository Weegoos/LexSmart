import io
import os
from datetime import date

from docx import Document

from app.schemas.contract import ContractCreate

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "templates")


def _format_date(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def _format_salary(amount: float, currency: str) -> str:
    formatted = f"{amount:,.0f}".replace(",", " ")
    return f"{formatted} {currency}"


# Mapping of placeholder tags to extraction lambdas
def _build_replacements(data: ContractCreate) -> dict[str, str]:
    return {
        "{{EMPLOYER_NAME}}": data.employer_name,
        "{{EMPLOYER_IIN_BIN}}": data.employer_iin_bin,
        "{{EMPLOYER_ADDRESS}}": data.employer_address,
        "{{EMPLOYEE_NAME}}": data.employee_name,
        "{{EMPLOYEE_IIN}}": data.employee_iin,
        "{{EMPLOYEE_ADDRESS}}": data.employee_address,
        "{{POSITION}}": data.position,
        "{{SALARY}}": _format_salary(data.salary, data.currency),
        "{{START_DATE}}": _format_date(data.start_date),
        "{{END_DATE}}": _format_date(data.end_date) if data.end_date else "бессрочный / indefinite",
        "{{PROBATION_MONTHS}}": str(data.probation_months),
        "{{WORK_SCHEDULE}}": data.work_schedule,
        "{{VACATION_DAYS}}": str(data.vacation_days),
        "{{CUSTOM_CLAUSES}}": data.custom_clauses or "",
        "{{CURRENT_DATE}}": _format_date(date.today()),
    }


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving formatting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(tag in full_text for tag in replacements):
        return

    for tag, value in replacements.items():
        if tag in full_text:
            full_text = full_text.replace(tag, value)

    # Rewrite runs: put all text in first run, clear the rest (preserves first run formatting)
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_contract_docx(data: ContractCreate) -> io.BytesIO:
    """Generate a .docx contract from template and return as in-memory buffer."""
    template_name = "labor_contract_ip.docx" if data.org_type == "IP" else "labor_contract_too.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    replacements = _build_replacements(data)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Also replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
